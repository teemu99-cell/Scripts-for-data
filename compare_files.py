#!/usr/bin/env python3
"""
compare_files.py
----------------
Compare two files — any mix of .txt, .docx, and .pdf — and show differences.

Dependencies for .docx / .pdf:
    pip install python-docx pymupdf --break-system-packages

Usage:
    python3 compare_files.py file1.txt file2.docx
    python3 compare_files.py report.pdf notes.txt -m side-by-side
    python3 compare_files.py a.docx b.pdf -m html
"""
import sys
import difflib
import argparse
from pathlib import Path

COLORS = {
    "green":  "\033[92m",
    "red":    "\033[91m",
    "blue":   "\033[94m",
    "yellow": "\033[93m",
    "reset":  "\033[0m",
    "bold":   "\033[1m",
}

def colorize(text, color):
    return f"{COLORS[color]}{text}{COLORS['reset']}"

# optional deps
try:
    import docx as _docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

def is_binary(path):
    try:
        with open(path, "rb") as f:
            return b"\x00" in f.read(8192)
    except OSError:
        return False

def extract_txt(path, encoding):
    encodings = [encoding] if encoding else ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.readlines(), enc
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode '{path}' — tried: {encodings}")

def extract_docx(path):
    if not HAS_DOCX:
        raise ImportError(
            "python-docx needed for .docx files.\n"
            "  pip install python-docx --break-system-packages"
        )
    doc = _docx.Document(str(path))
    lines = [p.text + "\n" for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text + "\n")
    return lines, "docx"

def extract_pdf(path):
    if not HAS_PDF:
        raise ImportError(
            "PyMuPDF needed for .pdf files.\n"
            "  pip install pymupdf --break-system-packages"
        )
    lines = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            for line in page.get_text().splitlines():
                if line.strip():
                    lines.append(line + "\n")
    return lines, "pdf"

def read_file(path, encoding=None):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if is_binary(path):
        raise ValueError(f"'{path.name}' appears to be a binary file.")
    return extract_txt(path, encoding)

def print_unified_diff(diff_lines):
    additions = deletions = 0
    for line in diff_lines:
        if line.startswith("+++") or line.startswith("---"):
            print(colorize(line, "bold"))
        elif line.startswith("+"):
            print(colorize(line, "green"))
            additions += 1
        elif line.startswith("-"):
            print(colorize(line, "red"))
            deletions += 1
        elif line.startswith("@@"):
            print(colorize(line, "blue"))
        else:
            print(line)
    return additions, deletions

def print_side_by_side(lines1, lines2, width=44):
    matcher = difflib.SequenceMatcher(None, lines1, lines2)
    additions = deletions = 0
    header = f"{'FILE 1':<{width}} | {'FILE 2':<{width}}"
    print(colorize(header, "bold"))
    print("-" * (width * 2 + 3))
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == "equal":
            for l, r in zip(lines1[i1:i2], lines2[j1:j2]):
                print(f"{l.rstrip():<{width}} | {r.rstrip()}")
        elif op == "replace":
            for l, r in zip(lines1[i1:i2], lines2[j1:j2]):
                print(colorize(f"{l.rstrip():<{width}}", "red") + " | " + colorize(r.rstrip(), "green"))
            deletions += i2 - i1
            additions += j2 - j1
        elif op == "delete":
            for l in lines1[i1:i2]:
                print(colorize(f"{l.rstrip():<{width}}", "red") + " | ")
            deletions += i2 - i1
        elif op == "insert":
            for r in lines2[j1:j2]:
                print(f"{'':<{width}} | " + colorize(r.rstrip(), "green"))
            additions += j2 - j1
    return additions, deletions

def compare_files(file1_path, file2_path, context_lines=3, mode="unified", encoding=None):
    p1, p2 = Path(file1_path), Path(file2_path)

    for p in (p1, p2):
        if not p.exists():
            print(colorize(f"Error: '{p}' does not exist.", "red"))
            sys.exit(1)
        if not p.is_file():
            print(colorize(f"Error: '{p}' is not a file.", "red"))
            sys.exit(1)
        if p.suffix.lower() not in {".txt", ".docx", ".pdf"}:
            print(colorize(f"Warning: '{p.name}' unsupported type '{p.suffix}'. "
                           "Supported: .txt .docx .pdf", "yellow"))

    try:
        lines1, fmt1 = read_file(p1, encoding)
        lines2, fmt2 = read_file(p2, encoding)
    except (ValueError, ImportError) as e:
        print(colorize(str(e), "red"))
        sys.exit(1)

    print(colorize(
        f"\nComparing: {p1.name} [{fmt1}]  vs  {p2.name} [{fmt2}]\n", "bold"
    ))

    additions = deletions = 0

    if mode == "unified":
        diff = list(difflib.unified_diff(
            lines1, lines2,
            fromfile=p1.name,
            tofile=p2.name,
            lineterm="",
            n=context_lines,
        ))
        if diff:
            print("Differences found:\n")
            additions, deletions = print_unified_diff(diff)
        else:
            print(colorize("Files are identical.", "green"))

    elif mode == "side-by-side":
        additions, deletions = print_side_by_side(lines1, lines2)

    elif mode == "html":
        html = difflib.HtmlDiff().make_file(
            lines1, lines2,
            fromdesc=f"{p1.name} [{fmt1}]",
            todesc=f"{p2.name} [{fmt2}]",
            context=True,
            numlines=context_lines,
        )
        out = Path("diff_output.html")
        out.write_text(html, encoding="utf-8")
        print(colorize(f"HTML diff written to '{out}'.", "green"))

    # summary
    print("\n" + colorize("=" * 50, "bold"))
    print(colorize("Summary:", "bold"))
    print(f"  {p1.name} [{fmt1}]: {colorize(str(len(lines1)), 'yellow')} lines")
    print(f"  {p2.name} [{fmt2}]: {colorize(str(len(lines2)), 'yellow')} lines")
    print(f"  Lines added:   {colorize(str(additions), 'green')}")
    print(f"  Lines deleted: {colorize(str(deletions), 'red')}")
    similarity = difflib.SequenceMatcher(None, lines1, lines2).ratio()
    sim_color = "green" if similarity >= 0.7 else "yellow" if similarity >= 0.4 else "red"
    print(f"  Similarity:    {colorize(f'{similarity:.1%}', sim_color)}")

def main():
    parser = argparse.ArgumentParser(
        description="Compare two files (.txt / .docx / .pdf) and highlight differences."
    )
    parser.add_argument("file1", help="First file  (.txt / .docx / .pdf)")
    parser.add_argument("file2", help="Second file (.txt / .docx / .pdf)")
    parser.add_argument("-c", "--context", type=int, default=3,
                        help="Context lines for unified diff (default: 3)")
    parser.add_argument("-m", "--mode", choices=["unified", "side-by-side", "html"],
                        default="unified", help="Output mode (default: unified)")
    parser.add_argument("-e", "--encoding",
                        help="Encoding for .txt files (default: auto-detect)")
    args = parser.parse_args()
    compare_files(args.file1, args.file2, args.context, args.mode, args.encoding)

if __name__ == "__main__":
    main()